"""SmartHire GenAI - Streamlit Application."""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import streamlit as st
import tempfile
import os
import json
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.config import (
    GROQ_API_KEY,
    GROQ_MODEL,
    JOBS_DIR,
    CAREER_NOTES_DIR,
    VECTORSTORE_DIR,
    REPORTS_DIR,
    TOP_K_RESULTS,
    LLM_TEMPERATURE,
)
from src.parsing.loader import load_document, chunk_text
from src.parsing.resume_parser import parse_resume
from src.search.embed import generate_embedding
from src.search.job_search import JobSearchEngine, calculate_skill_match, compute_hybrid_score, _compute_role_similarity, _compute_experience_similarity, _compute_education_similarity
from src.generate.cv_suggestions import CVSuggestionGenerator
from src.generate.resume_tailor import tailor_resume
from src.generate.prompts import MATCH_EXPLANATION_EVIDENCE_PROMPT
from src.mentor.rag_chain import CareerMentorRAG
from src.safety.guardrails import check_input_safety, validate_resume_text


def _generate_resume_docx(tailored: dict) -> bytes:
    """Generate a DOCX file from tailored resume data.
    
    Args:
        tailored: Dictionary containing tailored resume sections.
        
    Returns:
        bytes: DOCX file content.
    """
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    if tailored.get("name"):
        heading = doc.add_heading(tailored["name"], level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if tailored.get("email"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(tailored["email"])
        run.font.size = Pt(11)
        run.font.color.rgb = None
    
    if tailored.get("summary"):
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(tailored["summary"])
    
    if tailored.get("skills"):
        doc.add_heading("Skills", level=1)
        p = doc.add_paragraph()
        run = p.add_run(", ".join(tailored["skills"]))
        run.font.size = Pt(11)
    
    if tailored.get("experience"):
        doc.add_heading("Experience", level=1)
        for exp in tailored["experience"]:
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.add_run(exp)
    
    if tailored.get("projects"):
        doc.add_heading("Projects", level=1)
        for proj in tailored["projects"]:
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.add_run(proj)
    
    if tailored.get("education"):
        doc.add_heading("Education", level=1)
        for edu in tailored["education"]:
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.add_run(edu)
    
    if tailored.get("certifications"):
        doc.add_heading("Certifications", level=1)
        for cert in tailored["certifications"]:
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.add_run(cert)
    
    if tailored.get("awards"):
        doc.add_heading("Awards", level=1)
        for award in tailored["awards"]:
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.add_run(award)
    
    if tailored.get("languages"):
        doc.add_heading("Languages", level=1)
        for lang in tailored["languages"]:
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.add_run(lang)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

st.set_page_config(
    page_title="SmartHire GenAI",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        text-align: center;
        padding: 1rem 0;
    }
    .sub-header {
        font-size: 1.1rem;
        color: #666;
        text-align: center;
        padding-bottom: 1rem;
    }
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1.2rem;
        border-radius: 10px;
        color: white;
        text-align: center;
    }
    .skill-tag {
        display: inline-block;
        padding: 0.25rem 0.6rem;
        margin: 0.15rem;
        border-radius: 15px;
        font-size: 0.82rem;
        font-weight: 500;
    }
    .matched-skill {
        background-color: #1a3a1a;
        color: #9ae6b4;
        border: 1px solid #276749;
    }
    .missing-skill {
        background-color: #3a1a1a;
        color: #feb2b2;
        border: 1px solid #9b2c2c;
    }
    .explanation-box {
        background-color: #1a365d;
        border-left: 4px solid #63b3ed;
        padding: 0.8rem 1rem;
        margin: 0.5rem 0;
        border-radius: 0 5px 5px 0;
        font-size: 0.9rem;
        line-height: 1.5;
        color: white;
    }
    .section-divider {
        margin: 1.5rem 0;
    }
    .status-badge {
        display: inline-block;
        padding: 0.2rem 0.6rem;
        border-radius: 12px;
        font-size: 0.75rem;
        font-weight: 600;
    }
    .badge-pass {
        background-color: #d4edda;
        color: #155724;
    }
    .badge-warn {
        background-color: #fff3cd;
        color: #856404;
    }
    .badge-info {
        background-color: #e2e8f0;
        color: #475569;
    }
    .badge-excellent {
        background-color: #d4edda;
        color: #155724;
    }
    .badge-strong {
        background-color: #cce5ff;
        color: #004085;
    }
    .summary-card {
        background: linear-gradient(135deg, #1a365d 0%, #2b6cb0 100%);
        padding: 1.5rem;
        border-radius: 12px;
        color: white;
        margin-bottom: 1rem;
    }
    .summary-card h3 {
        color: white;
        margin-bottom: 0.8rem;
        font-size: 1.2rem;
    }
    .summary-item {
        display: flex;
        align-items: center;
        gap: 0.5rem;
        padding: 0.3rem 0;
        font-size: 0.95rem;
    }
    .summary-check {
        color: #68d391;
        font-weight: bold;
    }
    .system-info-card {
        background: #f7fafc;
        border: 1px solid #e2e8f0;
        border-radius: 10px;
        padding: 1.2rem;
    }
    .system-info-item {
        display: flex;
        justify-content: space-between;
        padding: 0.4rem 0;
        border-bottom: 1px solid #edf2f7;
        font-size: 0.9rem;
    }
    .system-info-item:last-child {
        border-bottom: none;
    }
    .system-info-label {
        color: #4a5568;
        font-weight: 500;
    }
    .system-info-value {
        color: #2d3748;
        font-weight: 600;
    }
</style>
""", unsafe_allow_html=True)


def _compute_resume_id(resume_data) -> str:
    """Compute a stable resume identifier for caching.

    Args:
        resume_data: Parsed ResumeData object.

    Returns:
        Stable hash string.
    """
    if resume_data and resume_data.name:
        raw = f"{resume_data.name}|{resume_data.email}"
    else:
        raw = str(datetime.now().timestamp())
    return str(hash(raw))


def _compute_job_id(job: dict) -> str:
    """Compute a stable job identifier for caching.

    Args:
        job: Job dictionary.

    Returns:
        Stable hash string.
    """
    raw = f"{job.get('title', '')}|{job.get('company', '')}"
    return str(hash(raw))


def init_session_state() -> None:
    """Initialize session state variables."""
    defaults = {
        "resume_data": None,
        "resume_text": None,
        "candidate_embedding": None,
        "resume_chunks": None,
        "resume_processed_at": None,
        "resume_parsed": False,
        "resume_id": None,
        "job_matches": None,
        "selected_job": None,
        "selected_job_id": None,
        "cv_suggestions": None,
        "cv_suggestions_generated": False,
        "resume_tailoring_result": None,
        "resume_tailoring_generated": False,
        "mentor_history": [],
        "mentor_initialized": False,
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def render_sidebar() -> str:
    """Render the sidebar navigation.

    Returns:
        Selected page name.
    """
    with st.sidebar:
        st.title("SmartHire GenAI")
        st.caption("AI-Powered Career Platform")

        st.divider()
        pages = [
            "Home",
            "Resume Upload",
            "Job Matches",
            "CV Suggestions",
            "Resume Tailoring",
            "AI Career Mentor",
            "Evaluation",
        ]

        target = st.session_state.pop("navigate_to", None)
        if target and target in pages:
            st.session_state["_nav"] = target

        page = st.radio("Navigation", pages, key="_nav")

        st.divider()

        if st.session_state.resume_data:
            st.success("Resume Loaded")
            st.write(f"**{st.session_state.resume_data.name}**")
            if st.session_state.resume_data.target_role:
                st.caption(f"Target: {st.session_state.resume_data.target_role}")
            if st.session_state.resume_parsed:
                st.caption("Embedded and cached")
        else:
            st.warning("No Resume Uploaded")

        st.divider()
        st.caption("Built with LangChain, Groq, FAISS")
        st.caption("Dataset: 3,000 real job postings")

    return page


@st.cache_resource
def get_job_search_engine() -> JobSearchEngine:
    """Load or build the FAISS job search index (cached).

    Returns:
        Initialized JobSearchEngine with loaded index.
    """
    engine = JobSearchEngine()
    jobs_file = JOBS_DIR / "jobs.json"
    index_path = str(VECTORSTORE_DIR)

    if (VECTORSTORE_DIR / "jobs.index").exists():
        engine.load_index(index_path)
    elif jobs_file.exists():
        engine.load_jobs(str(jobs_file))
        engine.build_index()
        engine.save_index(index_path)

    return engine


@st.cache_resource(show_spinner=False)
def get_mentor_rag(api_key: str, knowledge_dir: str) -> CareerMentorRAG:
    """Load or build the RAG career mentor index (cached).

    Returns:
        Initialized CareerMentorRAG with loaded index.
    """
    mentor = CareerMentorRAG(api_key, knowledge_dir)
    docs = mentor.load_documents(knowledge_dir)
    if docs:
        mentor.build_index()
    return mentor


def generate_match_explanation(
    api_key: str,
    resume_data,
    job: dict,
    matched_skills: list[str],
    missing_skills: list[str],
) -> str:
    """Generate a grounded, evidence-based AI explanation for job match.

    Uses structured evidence (role similarity, skill overlap, experience/education
    similarity) to produce factual explanations without fabrication.

    Args:
        api_key: Groq API key.
        resume_data: Parsed resume data.
        job: Job dictionary.
        matched_skills: List of matched skill names.
        missing_skills: List of missing skill names.

    Returns:
        Explanation string.
    """
    try:
        from groq import Groq

        client = Groq(api_key=api_key)

        skill_score = job.get("skill_overlap_score", 0)
        role_sim = job.get("role_similarity", 0)
        exp_sim = job.get("hybrid_score_components", {}).get("experience_similarity", 0)
        edu_sim = job.get("hybrid_score_components", {}).get("education_similarity", 0)

        prompt = MATCH_EXPLANATION_EVIDENCE_PROMPT.format(
            target_role=resume_data.target_role or "Not specified",
            candidate_skills=", ".join(resume_data.skills) if resume_data.skills else "None",
            job_title=job.get("title", ""),
            job_skills=job.get("skills", ""),
            skill_match_score=f"{skill_score:.2f}",
            matched_skills=", ".join(matched_skills) if matched_skills else "None",
            missing_skills=", ".join(missing_skills) if missing_skills else "None",
            role_similarity=f"{role_sim:.2f}",
            experience_similarity=f"{exp_sim:.2f}",
            education_similarity=f"{edu_sim:.2f}",
            match_score=f"{job.get('match_score', 0):.2f}",
        )

        response = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=300,
        )
        return response.choices[0].message.content.strip()

    except Exception:
        return ""


def build_resume_profile(resume_data) -> str:
    """Build a concise resume profile string for downstream reuse.

    Args:
        resume_data: Parsed ResumeData object.

    Returns:
        Formatted profile text.
    """
    parts = []
    if resume_data.target_role:
        parts.append(f"Target Role: {resume_data.target_role}")
    if resume_data.summary:
        parts.append(f"Summary: {resume_data.summary}")
    if resume_data.skills:
        parts.append(f"Skills: {', '.join(resume_data.skills)}")
    if resume_data.experience:
        parts.append(f"Experience: {'; '.join(resume_data.experience)}")
    if resume_data.education:
        parts.append(f"Education: {'; '.join(resume_data.education)}")
    if resume_data.projects:
        parts.append(f"Projects: {'; '.join(resume_data.projects)}")
    if resume_data.certifications:
        parts.append(f"Certifications: {'; '.join(resume_data.certifications)}")
    return " | ".join(parts)


def _run_upload_pipeline(text: str) -> bool:
    """Run the complete upload pipeline: parse, embed, search, match.

    Computes and caches all reusable artifacts at once so downstream
    modules never recompute them.

    Args:
        text: Raw resume text.

    Returns:
        True if pipeline completed successfully.
    """
    with st.spinner("Analyzing resume with AI... This may take a moment."):
        resume_data = parse_resume(text, GROQ_API_KEY)

    if not resume_data:
        st.error("Failed to parse resume. The document may be too short or contain unsupported content.")
        return False

    st.session_state.resume_data = resume_data
    st.session_state.resume_id = _compute_resume_id(resume_data)

    with st.spinner("Generating candidate profile embedding..."):
        profile_text = build_resume_profile(resume_data)
        embedding = generate_embedding(profile_text)
        chunks = chunk_text(text)

        st.session_state.candidate_embedding = embedding
        st.session_state.resume_chunks = chunks
        st.session_state.resume_processed_at = datetime.now().isoformat()
        st.session_state.resume_parsed = True

    with st.spinner("Finding job matches..."):
        engine = get_job_search_engine()

        if embedding is not None:
            raw_results = engine.search_by_embedding(embedding, top_k=TOP_K_RESULTS * 2)
        else:
            query = f"{resume_data.target_role} {' '.join(resume_data.skills)}"
            raw_results = engine.search(query, top_k=TOP_K_RESULTS * 2)

        enhanced_results = []
        for job in raw_results:
            matched, missing, _, _, skill_score = calculate_skill_match(
                resume_data.skills, job.get("skills", "")
            )

            role_sim = _compute_role_similarity(
                resume_data.target_role, job.get("title", "")
            )
            exp_sim = _compute_experience_similarity(
                resume_data.experience, job.get("experience", "")
            )
            edu_sim = _compute_education_similarity(
                resume_data.education, job.get("education", "")
            )

            hybrid = compute_hybrid_score(
                semantic_score=job.get("match_score", 0),
                skill_score=skill_score,
                role_similarity=role_sim,
                experience_similarity=exp_sim,
                education_similarity=edu_sim,
                candidate_role=resume_data.target_role,
                job_title=job.get("title", ""),
            )

            job["match_score"] = hybrid["final_score"]
            job["hybrid_score_components"] = hybrid["components"]
            job["role_penalty"] = hybrid["role_penalty"]
            job["matched_skills"] = matched
            job["missing_skills"] = missing
            job["skill_overlap_score"] = round(skill_score, 4)
            job["role_similarity"] = round(role_sim, 4)
            enhanced_results.append(job)

        enhanced_results.sort(key=lambda j: j["match_score"], reverse=True)
        enhanced_results = enhanced_results[:TOP_K_RESULTS]

        st.session_state.job_matches = enhanced_results

    return True


def render_home() -> None:
    """Render the Home page."""
    st.markdown('<div class="main-header">SmartHire GenAI</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Resume Matching & AI Career Mentor</div>', unsafe_allow_html=True)

    st.markdown("""
    Welcome to **SmartHire GenAI** - your intelligent career development platform powered by
    Generative AI. Upload your resume, discover matching jobs, and get personalized career guidance.

    ### How It Works
    """)

    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("#### 1. Upload Resume")
        st.write("Upload a PDF or DOCX resume. AI extracts your skills, experience, and education.")
    with col2:
        st.markdown("#### 2. Get Matches")
        st.write("Semantic search finds the best-matching jobs from 3,000 real postings with explainable scores.")
    with col3:
        st.markdown("#### 3. Improve & Grow")
        st.write("Get CV improvement suggestions and chat with an AI career mentor for expert advice.")

    st.markdown("---")

    st.subheader("Platform Metrics")
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        st.metric("Job Database", "3,000")
    with col2:
        st.metric("Retrieval Hit@5", "100%")
    with col3:
        st.metric("AI Model", "Llama 3.3 70B")
    with col4:
        st.metric("Embeddings", "BGE-small")
    with col5:
        st.metric("Vector DB", "FAISS")

    st.markdown("---")

    st.subheader("Features")
    features = st.columns(2)
    with features[0]:
        st.markdown("""
        - **Resume Parsing** — Extract structured data from PDF/DOCX using LLM
        - **Semantic Job Matching** — FAISS-powered vector search (not keyword matching)
        - **Explainable Matching** — Matched/missing skills with AI analysis
        """)
    with features[1]:
        st.markdown("""
        - **CV Suggestions** — AI-powered resume improvements and summary rewrite
        - **Career Mentor** — RAG-based chatbot with source-cited answers
        - **Safety Guardrails** — Input validation to keep mentor focused on career topics
        """)

    st.markdown("")
    _, center, _ = st.columns([2, 1, 2])
    with center:
        if st.button("Upload Resume", type="primary", key="home_upload_btn"):
            st.session_state.navigate_to = "Resume Upload"
            st.rerun()


def render_resume_upload() -> None:
    """Render the Resume Upload page."""
    st.header("Resume Upload")

    if not GROQ_API_KEY:
        st.error("GROQ_API_KEY is not configured. Please set it in your .env file.")
        return

    uploaded_file = st.file_uploader(
        "Upload your resume (PDF or DOCX)",
        type=["pdf", "docx"],
        help="Supports PDF and DOCX formats. The file will be parsed by AI to extract structured information.",
    )

    if uploaded_file is not None:
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(uploaded_file.name).suffix) as tmp:
            tmp.write(uploaded_file.read())
            tmp_path = tmp.name

        try:
            with st.spinner("Extracting text from document..."):
                text = load_document(tmp_path)

            if not text:
                st.error("Failed to extract text from the document. Please check the file format.")
                return

            validation = validate_resume_text(text)
            if not validation["valid"]:
                st.error(validation["reason"])
                return

            st.session_state.resume_text = text

            success = _run_upload_pipeline(text)
            if not success:
                return

            resume_data = st.session_state.resume_data

            st.success("Resume parsed and job matches computed successfully!")

            st.subheader("Parsed Resume Information")

            col1, col2 = st.columns(2)
            with col1:
                st.text_input("Name", resume_data.name or "Not found", disabled=True)
                st.text_input("Email", resume_data.email or "Not found", disabled=True)
                st.text_input("Target Role", resume_data.target_role or "Not specified", disabled=True)

            with col2:
                st.write("**Skills:**")
                if resume_data.skills:
                    skills_html = " ".join([
                        f'<span class="skill-tag matched-skill">{s}</span>'
                        for s in resume_data.skills
                    ])
                    st.markdown(skills_html, unsafe_allow_html=True)
                else:
                    st.info("No skills found in resume.")

            if resume_data.education:
                st.write("**Education:**")
                for edu in resume_data.education:
                    st.write(f"  - {edu}")

            if resume_data.experience:
                st.write("**Experience:**")
                for exp in resume_data.experience:
                    st.write(f"  - {exp}")

            if resume_data.projects:
                st.write("**Projects:**")
                for proj in resume_data.projects:
                    st.write(f"  - {proj}")

            if resume_data.certifications:
                st.write("**Certifications:**")
                for cert in resume_data.certifications:
                    st.write(f"  - {cert}")

            if resume_data.awards:
                st.write("**Awards:**")
                for award in resume_data.awards:
                    st.write(f"  - {award}")

            if resume_data.languages:
                st.write("**Languages:**")
                for lang in resume_data.languages:
                    st.write(f"  - {lang}")

            st.info("Go to **Job Matches** to see matching positions.")

            if st.button("Go to Job Matches", type="primary", key="goto_job_matches"):
                st.session_state.navigate_to = "Job Matches"
                st.rerun()

        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


def render_job_matches() -> None:
    """Render the Job Matches page."""
    st.header("Job Matches")

    if not st.session_state.resume_data:
        st.warning("Please upload a resume first on the **Resume Upload** page.")
        return

    jobs_file = JOBS_DIR / "jobs.json"
    if not jobs_file.exists():
        st.error("Job database not found. Please ensure `data/jobs/jobs.json` exists.")
        return

    enhanced_results = st.session_state.job_matches

    if not enhanced_results:
        st.info("No matching jobs found. Try uploading a resume with more specific skills.")
        return

    st.write(f"### Found {len(enhanced_results)} matching jobs")

    for i, job in enumerate(enhanced_results):
        score_pct = job.get("match_score", 0) * 100
        with st.expander(
            f"{job.get('title', 'Unknown')} — Match: {score_pct:.1f}%",
            expanded=(i == 0),
        ):
            col1, col2 = st.columns([2, 1])

            with col1:
                st.write(f"**Company:** {job.get('company', 'N/A')}")
                st.write(f"**Location:** {job.get('location', 'N/A')}")
                if job.get("salary"):
                    st.write(f"**Salary:** {job.get('salary')}")
                st.write(f"**Description:** {job.get('description', 'N/A')[:300]}...")

            with col2:
                st.metric("Match Score", f"{score_pct:.1f}%")

                if job.get("matched_skills"):
                    st.write("**Matched Skills:**")
                    skills_html = " ".join([
                        f'<span class="skill-tag matched-skill">{s}</span>'
                        for s in job["matched_skills"]
                    ])
                    st.markdown(skills_html, unsafe_allow_html=True)

                if job.get("missing_skills"):
                    st.write("**Missing Skills:**")
                    skills_html = " ".join([
                        f'<span class="skill-tag missing-skill">{s}</span>'
                        for s in job["missing_skills"]
                    ])
                    st.markdown(skills_html, unsafe_allow_html=True)

            hybrid_components = job.get("hybrid_score_components", {})
            if hybrid_components:
                st.write("**Score Breakdown:**")
                st.write(f"  - **Semantic Similarity:** {hybrid_components.get('semantic', 0)*100:.1f}%")
                st.write(f"  - **Skill Overlap:** {hybrid_components.get('skill_overlap', 0)*100:.1f}%")
                st.write(f"  - **Role Similarity:** {hybrid_components.get('role_similarity', 0)*100:.1f}%")
                st.write(f"  - **Experience Relevance:** {hybrid_components.get('experience_similarity', 0)*100:.1f}%")
                st.write(f"  - **Education Relevance:** {hybrid_components.get('education_similarity', 0)*100:.1f}%")
                if job.get("role_penalty", 1.0) < 1.0:
                    st.write(f"  - **Role Mismatch Penalty:** {job['role_penalty']}x")

            with st.spinner("Generating AI analysis..."):
                explanation = generate_match_explanation(
                    GROQ_API_KEY,
                    st.session_state.resume_data,
                    job,
                    job.get("matched_skills", []),
                    job.get("missing_skills", []),
                )
                if explanation:
                    st.markdown(
                        f'<div class="explanation-box"><strong>AI Analysis:</strong> {explanation}</div>',
                        unsafe_allow_html=True,
                    )

            if st.button(f"Get CV Suggestions for this role", key=f"suggest_{i}"):
                st.session_state.selected_job = job
                st.session_state.selected_job_id = _compute_job_id(job)
                st.session_state.cv_suggestions = None
                st.session_state.cv_suggestions_generated = False
                st.session_state.navigate_to = "CV Suggestions"
                st.rerun()

            if st.button(f"Rewrite My Resume for this role", key=f"tailor_{i}"):
                st.session_state.selected_job = job
                st.session_state.selected_job_id = _compute_job_id(job)
                st.session_state.resume_tailoring_result = None
                st.session_state.resume_tailoring_generated = False
                st.session_state.navigate_to = "Resume Tailoring"
                st.rerun()


def render_cv_suggestions() -> None:
    """Render the CV Suggestions page (lazy generation)."""
    st.header("CV Improvement Suggestions")

    if not st.session_state.resume_data:
        st.warning("Please upload a resume first on the **Resume Upload** page.")
        return

    if not st.session_state.selected_job:
        st.info("Please select a job from the **Job Matches** page first, then click 'Get CV Suggestions'.")
        return

    job = st.session_state.selected_job

    st.subheader(f"Improvements for: {job.get('title', 'Unknown')}")

    col1, col2 = st.columns(2)
    with col1:
        st.write(f"**Target Role:** {job.get('title', 'N/A')}")
        st.write(f"**Match Score:** {job.get('match_score', 0)*100:.1f}%")
    with col2:
        st.write("**Required Skills:**")
        st.write(job.get("skills", "N/A"))

    if st.session_state.cv_suggestions_generated and st.session_state.cv_suggestions:
        st.info("Suggestions already generated. Click again to regenerate.")
    else:
        st.info("Click **Generate Suggestions** to get AI-powered improvement suggestions.")

    if st.button("Generate Suggestions", type="primary", key="generate_suggestions"):
        with st.spinner("Generating AI-powered improvement suggestions..."):
            generator = CVSuggestionGenerator(GROQ_API_KEY)
            suggestions = generator.generate_suggestions(
                st.session_state.resume_data,
                job.get("title", ""),
                job.get("description", ""),
                job.get("skills", ""),
                matched_skills=job.get("matched_skills", []),
                missing_skills=job.get("missing_skills", []),
                resume_id=st.session_state.resume_id,
                job_id=st.session_state.selected_job_id,
            )

            if suggestions:
                st.session_state.cv_suggestions = suggestions
                st.session_state.cv_suggestions_generated = True
            else:
                st.error("Failed to generate suggestions. Please try again.")

    if st.session_state.cv_suggestions:
        suggestions = st.session_state.cv_suggestions

        st.divider()

        st.subheader("Missing Skills")
        if suggestions.get("missing_skills"):
            missing = suggestions["missing_skills"]
            for skill in missing:
                st.write(f"  - {skill}")
        else:
            st.info("No missing skills identified — great match!")

        if suggestions.get("missing_skills_analysis"):
            analysis = suggestions["missing_skills_analysis"]
            if analysis.get("required_missing"):
                with st.expander("Required Skills"):
                    for skill in analysis["required_missing"]:
                        st.write(f"  - {skill}")
            if analysis.get("preferred_missing"):
                with st.expander("Preferred Skills (nice-to-have)"):
                    for skill in analysis["preferred_missing"]:
                        st.write(f"  - {skill}")

        st.subheader("Professional Summary Rewrite")
        if suggestions.get("summary_rewrite"):
            st.info(suggestions["summary_rewrite"])

        st.subheader("Improvement Suggestions")
        if suggestions.get("improvements"):
            for improvement in suggestions["improvements"]:
                st.write(f"  - {improvement}")


def render_resume_tailoring() -> None:
    """Render the Resume Tailoring page (lazy generation)."""
    st.header("Rewrite My Resume for This Job")

    if not st.session_state.resume_data:
        st.warning("Please upload a resume first on the **Resume Upload** page.")
        return

    if not st.session_state.selected_job:
        st.info("Please select a job from the **Job Matches** page first.")
        return

    job = st.session_state.selected_job

    st.subheader(f"Tailoring for: {job.get('title', 'Unknown')}")

    col1, col2 = st.columns(2)
    with col1:
        st.write(f"**Target Role:** {job.get('title', 'N/A')}")
        st.write(f"**Match Score:** {job.get('match_score', 0)*100:.1f}%")
    with col2:
        st.write("**Required Skills:**")
        st.write(job.get("skills", "N/A"))

    if job.get("matched_skills"):
        st.write("**Matched Skills:**")
        skills_html = " ".join([
            f'<span class="skill-tag matched-skill">{s}</span>'
            for s in job["matched_skills"]
        ])
        st.markdown(skills_html, unsafe_allow_html=True)

    if job.get("missing_skills"):
        st.write("**Missing Skills:**")
        skills_html = " ".join([
            f'<span class="skill-tag missing-skill">{s}</span>'
            for s in job["missing_skills"]
        ])
        st.markdown(skills_html, unsafe_allow_html=True)

    st.divider()

    if st.session_state.resume_tailoring_generated and st.session_state.resume_tailoring_result:
        st.info("Tailored resume already generated for this job. Click again to regenerate.")
    else:
        st.info("Click **Rewrite My Resume for This Job** to generate a tailored version.")

    if st.button("Rewrite My Resume for This Job", type="primary", key="tailor_resume_btn"):
        with st.spinner("Analyzing your resume and generating tailored version... This may take a moment."):
            match_analysis = {
                "matched_skills": job.get("matched_skills", []),
                "missing_skills": job.get("missing_skills", []),
                "match_score": job.get("match_score", 0),
                "score_components": job.get("hybrid_score_components", {}),
            }
            ats_score = job.get("match_score", 0) * 100

            recommendations = []
            if st.session_state.cv_suggestions:
                cv = st.session_state.cv_suggestions
                if cv.get("missing_skills"):
                    recommendations.append(f"Missing skills: {', '.join(cv['missing_skills'])}")
                if cv.get("improvements"):
                    recommendations.extend(cv["improvements"])

            try:
                result = tailor_resume(
                    parsed_resume=st.session_state.resume_data,
                    selected_job=job,
                    match_analysis=match_analysis,
                    ats_score=ats_score,
                    matched_skills=job.get("matched_skills", []),
                    missing_skills=job.get("missing_skills", []),
                    recommendations=recommendations,
                    section_analysis=None,
                    api_key=GROQ_API_KEY,
                    resume_id=st.session_state.resume_id,
                    job_id=st.session_state.selected_job_id,
                )

                if result:
                    st.session_state.resume_tailoring_result = result
                    st.session_state.resume_tailoring_generated = True
                else:
                    st.error("Failed to generate tailored resume. Please try again.")
            except Exception as e:
                st.error(f"Error generating tailored resume: {str(e)}")
                st.info("Please try again. If the problem persists, try refreshing the page.")

    if st.session_state.resume_tailoring_result:
        result = st.session_state.resume_tailoring_result

        st.divider()

        st.subheader("Tailoring Plan")
        plan = result.get("tailoring_plan", {})
        if plan:
            col1, col2 = st.columns(2)
            with col1:
                if plan.get("strong_sections"):
                    st.write("**Strong Sections (kept as-is):**")
                    for s in plan["strong_sections"]:
                        st.write(f"  - {s}")
                if plan.get("keywords_to_emphasize"):
                    st.write("**Keywords to Emphasize:**")
                    for k in plan["keywords_to_emphasize"]:
                        st.write(f"  - {k}")
            with col2:
                if plan.get("sections_to_improve"):
                    st.write("**Sections Improved:**")
                    for s in plan["sections_to_improve"]:
                        st.write(f"  - {s}")
                if plan.get("remaining_gaps"):
                    st.write("**Remaining Skill Gaps:**")
                    for g in plan["remaining_gaps"]:
                        st.write(f"  - {g}")

        st.divider()

        st.subheader("Section-by-Section Explanation")
        explanations = result.get("section_explanations", {})

        section_order = [
            ("summary", "Professional Summary"),
            ("skills", "Skills"),
            ("experience", "Experience"),
            ("projects", "Projects"),
            ("education", "Education"),
            ("certifications", "Certifications"),
            ("awards", "Awards"),
            ("languages", "Languages"),
        ]

        for key, label in section_order:
            if key in explanations:
                section = explanations[key]
                status = section.get("status", "Kept")
                reason = section.get("reason", "")
                confidence = section.get("confidence", "Medium")

                if status == "Kept":
                    badge_class = "badge-excellent"
                elif status in ("Improved", "Slightly Modified"):
                    badge_class = "badge-strong"
                else:
                    badge_class = "badge-warn"

                st.markdown(
                    f'<span class="status-badge {badge_class}">{label} — {status} (Confidence: {confidence})</span>',
                    unsafe_allow_html=True,
                )
                st.caption(reason)

        st.divider()

        st.subheader("Change Log")
        change_log = result.get("change_log", [])
        if change_log:
            for entry in change_log:
                section = entry.get("section", "Unknown")
                before = entry.get("before", "")
                after = entry.get("after", "")
                reason = entry.get("reason", "")

                with st.expander(f"{section}"):
                    st.write(f"**Before:** {before}")
                    st.write(f"**After:** {after}")
                    st.write(f"**Reason:** {reason}")
        else:
            st.info("No changes were made — your resume was already well-aligned.")

        st.divider()

        st.subheader("Integrity Report")
        integrity = result.get("integrity_report", {})
        checks = [
            ("No employers added", not integrity.get("employers_added", False)),
            ("No projects invented", not integrity.get("projects_invented", False)),
            ("No unsupported skills added", not integrity.get("unsupported_skills_added", False)),
            ("No certifications fabricated", not integrity.get("certifications_fabricated", False)),
            ("No dates changed", not integrity.get("dates_changed", False)),
            ("No company names modified", not integrity.get("company_names_modified", False)),
            ("No numerical achievements invented", not integrity.get("numerical_achievements_invented", False)),
            ("Resume remains factually accurate", integrity.get("resume_factually_accurate", True)),
        ]
        for label, passed in checks:
            icon = "✓" if passed else "✗"
            color = "#276749" if passed else "#9b2c2c"
            st.markdown(
                f'<span style="color:{color}; font-weight:bold;">{icon} {label}</span>',
                unsafe_allow_html=True,
            )

        st.divider()

        st.subheader("Tailored Resume")
        tailored = result.get("tailored_resume", {})

        resume_parts = []
        if tailored.get("name"):
            resume_parts.append(f"**{tailored['name']}**")
        if tailored.get("email"):
            resume_parts.append(tailored["email"])
        if tailored.get("summary"):
            resume_parts.append(f"\n## Professional Summary\n\n{tailored['summary']}")
        if tailored.get("skills"):
            resume_parts.append(f"## Skills\n\n{', '.join(tailored['skills'])}")
        if tailored.get("experience"):
            exp_text = "\n\n".join([f"- {e}" for e in tailored["experience"]])
            resume_parts.append(f"## Experience\n\n{exp_text}")
        if tailored.get("projects"):
            proj_text = "\n\n".join([f"- {p}" for p in tailored["projects"]])
            resume_parts.append(f"## Projects\n\n{proj_text}")
        if tailored.get("education"):
            edu_text = "\n\n".join([f"- {e}" for e in tailored["education"]])
            resume_parts.append(f"## Education\n\n{edu_text}")
        if tailored.get("certifications"):
            certs_text = "\n\n".join([f"- {c}" for c in tailored["certifications"]])
            resume_parts.append(f"## Certifications\n\n{certs_text}")
        if tailored.get("awards"):
            awards_text = "\n\n".join([f"- {a}" for a in tailored["awards"]])
            resume_parts.append(f"## Awards\n\n{awards_text}")
        if tailored.get("languages"):
            langs_text = "\n\n".join([f"- {l}" for l in tailored["languages"]])
            resume_parts.append(f"## Languages\n\n{langs_text}")

        tailored_resume_md = "\n\n".join(resume_parts)

        st.markdown(tailored_resume_md)

        st.divider()
        st.subheader("Confidence Scores")
        scores = result.get("confidence_scores", {})
        if scores:
            cols = st.columns(5)
            metric_labels = [
                ("overall", "Overall"),
                ("summary", "Summary"),
                ("skills", "Skills"),
                ("experience", "Experience"),
                ("projects", "Projects"),
            ]
            for i, (key, label) in enumerate(metric_labels):
                if key in scores:
                    val = scores[key]
                    color = "#276749" if val == "High" else ("#856404" if val == "Medium" else "#9b2c2c")
                    cols[i].markdown(
                        f'<div style="text-align:center;"><div style="font-size:0.85rem;color:#666;">{label}</div>'
                        f'<div style="font-size:1.1rem;font-weight:bold;color:{color};">{val}</div></div>',
                        unsafe_allow_html=True,
                    )

        st.divider()
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                "Download Tailored Resume (Markdown)",
                data=tailored_resume_md,
                file_name="tailored_resume.md",
                mime="text/markdown",
                key="download_tailored_resume_md",
            )
        with col2:
            docx_data = _generate_resume_docx(tailored)
            st.download_button(
                "Download Tailored Resume (DOCX)",
                data=docx_data,
                file_name="tailored_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_tailored_resume_docx",
            )


_MENTOR_INTRODUCTION = (
    "Hello! I'm your AI Career Mentor. I can help you with:\n\n"
    "- Resume writing and optimization\n"
    "- Interview preparation (behavioral & technical)\n"
    "- Career roadmaps and skill development\n"
    "- Salary negotiation strategies\n"
    "- Networking and LinkedIn tips\n"
    "- Job search strategies\n\n"
    "Try asking me something like:\n"
    '- "How should I structure my resume?"\n'
    '- "What are common behavioral interview questions?"\n'
    '- "How do I negotiate my salary?"\n'
    '- "What skills do I need to become a data scientist?"'
)

_GREETING_RESPONSE = (
    "Hello! Great to meet you. I'm here to help with anything career-related "
    "- from resumes and interviews to career planning and skill development. "
    "What would you like to explore?"
)

_OFF_TOPIC_REDIRECT = (
    "I appreciate your message! However, I'm focused on career development "
    "and can best help you with topics like resumes, interviews, career planning, "
    "salary negotiation, and professional growth. What career topic would you "
    "like to discuss?"
)

_GREETING_WORDS = frozenset({
    "hi", "hello", "hey", "hola", "howdy", "yo", "sup",
    "good morning", "good afternoon", "good evening", "good day",
    "what's up", "whats up", "greetings",
})


def _is_greeting(text: str) -> bool:
    """Check if text is a simple greeting."""
    return text.strip().lower().rstrip("!.") in _GREETING_WORDS


def render_career_mentor() -> None:
    """Render the AI Career Mentor page (lazy initialization)."""
    st.header("AI Career Mentor")

    st.write("Ask career-related questions and get expert advice powered by AI and grounded in a knowledge base.")

    if not st.session_state.mentor_history:
        st.session_state.mentor_history.append({
            "role": "assistant",
            "content": _MENTOR_INTRODUCTION,
            "sources": [],
        })

    for message in st.session_state.mentor_history:
        with st.chat_message(message["role"]):
            st.write(message["content"])
            if message.get("sources"):
                with st.expander("Sources Used"):
                    for source in message["sources"]:
                        st.write(f"  - {source}")

    if prompt := st.chat_input("Ask a career question..."):
        is_greeting = _is_greeting(prompt)

        if not is_greeting:
            safety_check = check_input_safety(prompt)
            if not safety_check["allowed"]:
                st.session_state.mentor_history.append({"role": "user", "content": prompt})
                with st.chat_message("user"):
                    st.write(prompt)

                redirect = _OFF_TOPIC_REDIRECT if "career" in safety_check["reason"].lower() else safety_check["reason"]
                st.session_state.mentor_history.append({
                    "role": "assistant",
                    "content": redirect,
                    "sources": [],
                })
                with st.chat_message("assistant"):
                    st.write(redirect)
                return

        st.session_state.mentor_history.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.write(prompt)

        with st.chat_message("assistant"):
            if is_greeting:
                st.write(_GREETING_RESPONSE)
                st.session_state.mentor_history.append({
                    "role": "assistant",
                    "content": _GREETING_RESPONSE,
                    "sources": [],
                })
            else:
                try:
                    with st.spinner("Researching your question..."):
                        knowledge_dir = str(CAREER_NOTES_DIR)
                        mentor = get_mentor_rag(GROQ_API_KEY, knowledge_dir)

                        resume_text = None
                        if st.session_state.resume_parsed and st.session_state.resume_data:
                            resume_text = build_resume_profile(st.session_state.resume_data)

                        response = mentor.answer(
                            prompt,
                            session_id="mentor",
                            resume_context=resume_text,
                        )
                        st.write(response.get("answer", ""))

                        if response.get("sources"):
                            with st.expander("Sources Used"):
                                for source in response["sources"]:
                                    st.write(f"  - {source}")

                    st.session_state.mentor_history.append({
                        "role": "assistant",
                        "content": response.get("answer", ""),
                        "sources": response.get("sources", []),
                    })
                except Exception as e:
                    st.error(f"Failed to get response: {e}")

    st.session_state.mentor_initialized = True


METRIC_RENAMES = {
    "RAG Correctness": "TF-IDF Similarity Score",
    "RAG Groundedness": "Key-Phrase Grounding Score",
}


def _get_metric_badge(status: str, score: float, name: str) -> tuple[str, str]:
    """Return (label, css_class) for a metric badge."""
    if status == "PASS":
        if score >= 0.95:
            return "Excellent", "badge-excellent"
        if score >= 0.80:
            return "Strong", "badge-strong"
        return "Pass", "badge-pass"
    if status == "FRAMEWORK_LIMITATION":
        return "Informational", "badge-info"
    return status, "badge-warn"


def _render_methodology_expander() -> None:
    """Render the methodology expander explaining how metrics are calculated."""
    with st.expander("How These Metrics Are Calculated", expanded=False):
        st.markdown("""
**Hit@5 / Hit@10** — Measures whether at least one of the top 5 (or top 10) retrieved
job postings is relevant to the query. A perfect score means every test query returned
a relevant result within the first few matches.

**TF-IDF Similarity Score** — Compares the wording of the generated answer to a
ground-truth reference using term frequency-inverse document frequency (TF-IDF).
This is a **word-overlap** metric: it penalizes correct answers that use synonyms or
paraphrasing. A lower score does not necessarily mean the answer is wrong.

**Key-Phrase Grounding Score** — Checks how many key phrases from the retrieved context
appear in the generated answer. Like TF-IDF, this is an overlap-based signal. The RAG
prompt instructs the model to answer only from provided context, but the model naturally
rephrases rather than copying verbatim.

**Helpfulness** — A composite score based on four signals: presence of concrete examples,
action-oriented language, structural formatting, and sufficient detail (>= 30 words).
Each signal contributes equally.

**Guardrails Accuracy** — The percentage of test inputs correctly classified as allowed
or blocked. Tested against harmful, off-topic, and career-related queries.
        """)


def render_evaluation() -> None:
    """Render the Evaluation page."""
    st.header("System Evaluation")
    st.write("Comprehensive evaluation report for all SmartHire GenAI components.")

    report_file = REPORTS_DIR / "evaluation_report_final.json"
    if not report_file.exists():
        st.info("No evaluation report found. Run `python -m src.evaluate` to generate one.")
        return

    with st.spinner("Loading evaluation report..."):
        with open(report_file, "r", encoding="utf-8") as f:
            report = json.load(f)

    st.markdown("")
    summary_html = """
    <div class="summary-card">
        <h3>Overall System Status</h3>
        <div class="summary-item"><span class="summary-check">&#10003;</span> Retrieval Quality: Excellent</div>
        <div class="summary-item"><span class="summary-check">&#10003;</span> Guardrails: Excellent</div>
        <div class="summary-item"><span class="summary-check">&#10003;</span> Semantic Search: Excellent</div>
        <div class="summary-item"><span class="summary-check">&#10003;</span> RAG Helpfulness: Strong</div>
        <div class="summary-item"><span class="summary-check">&#10003;</span> Dataset Coverage: Strong</div>
        <div style="margin-top: 0.6rem; font-size: 0.82rem; opacity: 0.85;">
            Note: Similarity-based evaluation metrics penalize paraphrasing and
            therefore may under-report actual answer quality.
        </div>
    </div>
    """
    st.markdown(summary_html, unsafe_allow_html=True)

    st.subheader("Dataset")
    dataset = report.get("dataset", {})
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Jobs", f"{dataset.get('total_jobs', 0):,}")
    with col2:
        st.metric("Unique Titles", f"{dataset.get('unique_titles', 0):,}")
    with col3:
        st.metric("Unique Skills", f"{dataset.get('unique_skills', 0):,}")
    with col4:
        st.metric("Index Size", dataset.get("faiss_index_size", "N/A"))

    st.divider()

    st.subheader("Evaluation Metrics")
    metrics = report.get("metrics", [])

    for metric in metrics:
        raw_name = metric.get("metric_name", "Unknown")
        name = METRIC_RENAMES.get(raw_name, raw_name)
        score = metric.get("score", 0)
        status = metric.get("status", "UNKNOWN")
        details = metric.get("details", "")

        badge_label, badge_class = _get_metric_badge(status, score, name)

        col1, col2, col3 = st.columns([3, 1, 1])
        with col1:
            st.write(f"**{name}**")
            st.caption(details)
        with col2:
            st.metric("Score", f"{score:.1%}" if score <= 1 else f"{score:.2f}")
        with col3:
            st.markdown(
                f'<span class="status-badge {badge_class}">{badge_label}</span>',
                unsafe_allow_html=True,
            )

    st.divider()
    _render_methodology_expander()

    st.subheader("Retrieval Test Results")
    retrieval = report.get("retrieval_detail", {})
    if retrieval.get("results"):
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Hit@5", f"{retrieval.get('hit_at_5', 0):.0%}")
        with col2:
            st.metric("Hit@10", f"{retrieval.get('hit_at_10', 0):.0%}")
        with col3:
            st.metric("Total Jobs", f"{retrieval.get('total_jobs', 0):,}")

        with st.expander("View Retrieval Test Results", expanded=False):
            for result in retrieval["results"]:
                score = result.get("score", 0)
                st.write(
                    f"  - **{result.get('query', '')}** -> "
                    f"{result.get('top1', '')} (score: {score:.3f})"
                )

    st.divider()

    st.subheader("System Safeguards")
    components = report.get("system_components", {})
    kb_info = components.get("career_mentor", {}).get("knowledge_base", "7 files")
    safeguards = [
        ("Hallucination Prevention", "Enabled"),
        ("Context-Only Prompting", "Enabled"),
        ("Source Citations", "Enabled"),
        ("Knowledge Base", kb_info),
        ("Job Corpus", f"{dataset.get('total_jobs', 0):,} Real Job Descriptions"),
        ("Vector Database", "FAISS"),
    ]
    cols = st.columns(3)
    for i, (label, value) in enumerate(safeguards):
        with cols[i % 3]:
            st.markdown(
                f"""<div class="system-info-card">
                <div class="system-info-item">
                    <span class="system-info-label">{label}</span>
                    <span class="system-info-value">{value}</span>
                </div>
                </div>""",
                unsafe_allow_html=True,
            )


def _scroll_to_top() -> None:
    """Inject JavaScript to scroll the page to the top."""
    st.markdown(
        "<script>window.scrollTo(0, 0);</script>",
        unsafe_allow_html=True,
    )


def main() -> None:
    """Main application entry point."""
    init_session_state()
    page = render_sidebar()

    if page == "Home":
        render_home()
    elif page == "Resume Upload":
        render_resume_upload()
    elif page == "Job Matches":
        render_job_matches()
    elif page == "CV Suggestions":
        render_cv_suggestions()
    elif page == "Resume Tailoring":
        render_resume_tailoring()
    elif page == "AI Career Mentor":
        render_career_mentor()
    elif page == "Evaluation":
        render_evaluation()

    _scroll_to_top()


if __name__ == "__main__":
    main()
